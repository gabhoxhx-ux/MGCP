import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Configuración
INPUT_FILE = 'TESTING_PLAN.md'
OUTPUT_FILE = 'Guia_QA_MGCP.docx'

def aplicar_formato_markdown(paragraph, text):
    """
    Procesa negritas (**texto**) y código en línea (`texto`).
    """
    # Dividir por negritas y código
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0) # Negro fuerte
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.color.rgb = RGBColor(200, 50, 50) # Rojo oscuro para código inline
        else:
            if part:
                paragraph.add_run(part)

def agregar_bloque_codigo(doc, texto_codigo):
    """
    Crea un párrafo con estilo de bloque de código (fondo grisáceo simulado o fuente mono).
    """
    p = doc.add_paragraph()
    runner = p.add_run(texto_codigo)
    runner.font.name = 'Consolas'
    runner.font.size = Pt(9)
    runner.font.color.rgb = RGBColor(30, 30, 30)
    
    # Indentación para que destaque
    p.paragraph_format.left_indent = Inches(0.5)
    
    # Añadir un borde simple (opcional, requiere manipulación XML, aquí usamos solo fondo/indentación simple)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)

def main():
    if not os.path.exists(INPUT_FILE):
        print(f"ERROR: No se encuentra '{INPUT_FILE}'. Asegúrate de que está en la carpeta.")
        return

    doc = Document()
    
    # --- Estilos Globales ---
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    print(f"Procesando {INPUT_FILE}...")

    with open(INPUT_FILE, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    en_bloque_codigo = False
    
    for line in lines:
        stripped = line.strip()
        
        # 1. Detección de Bloques de Código (```)
        if stripped.startswith('```'):
            en_bloque_codigo = not en_bloque_codigo
            continue # Saltamos la línea de los backticks
        
        if en_bloque_codigo:
            agregar_bloque_codigo(doc, line.rstrip()) # Mantenemos indentación interna
            continue

        # Si línea vacía fuera de código, saltar
        if not stripped:
            continue

        # 2. Encabezados (Títulos)
        if stripped.startswith('# '):
            # Título Principal
            head = doc.add_heading(stripped[2:], level=0)
            head.alignment = WD_ALIGN_PARAGRAPH.CENTER
            head.runs[0].font.color.rgb = RGBColor(0, 51, 102) # Azul oscuro profesional
        
        elif stripped.startswith('## '):
            head = doc.add_heading(stripped[3:], level=1)
            head.runs[0].font.color.rgb = RGBColor(46, 116, 181) # Azul medio
        
        elif stripped.startswith('### '):
            head = doc.add_heading(stripped[4:], level=2)
            head.runs[0].font.color.rgb = RGBColor(68, 84, 106) # Gris azulado
        
        elif stripped.startswith('#### '):
            # Nivel 4 lo hacemos como texto negrita subrayado o heading 3 pequeño
            p = doc.add_heading(stripped[5:], level=3)
            p.runs[0].font.size = Pt(12)
            p.runs[0].font.italic = True

        # 3. Listas de Verificación (Checkboxes - [ ])
        elif stripped.startswith('- [ ]'):
            p = doc.add_paragraph()
            # Insertar símbolo de checkbox manual
            run_check = p.add_run('☐ ') 
            run_check.font.size = Pt(14)
            run_check.font.bold = True
            
            texto_tarea = stripped[5:]
            aplicar_formato_markdown(p, texto_tarea)

        # 4. Listas normales (bullets)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            texto_lista = stripped[2:]
            aplicar_formato_markdown(p, texto_lista)

        # 5. Párrafos normales
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            aplicar_formato_markdown(p, stripped)

    # Guardar
    doc.save(OUTPUT_FILE)
    print(f"¡Listo! Informe generado: {OUTPUT_FILE}")

if __name__ == "__main__":
    main()